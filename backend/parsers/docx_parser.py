from pathlib import Path

from docx import Document


def parse_docx(file_path: str | Path) -> str:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {path}")
    if not path.suffix.lower() == ".docx":
        raise ValueError(f"File is not a DOCX: {path}")

    try:
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs).strip()
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {e}")


def parse_docx_bytes(data: bytes) -> str:
    import io

    try:
        doc = Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs).strip()
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX bytes: {e}")
