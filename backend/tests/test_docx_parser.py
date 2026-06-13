import pytest
from pathlib import Path
from parsers.docx_parser import parse_docx, parse_docx_bytes


def test_parse_docx_not_found():
    with pytest.raises(FileNotFoundError):
        parse_docx("nonexistent.docx")


def test_parse_docx_wrong_extension(tmp_path):
    txt_file = tmp_path / "test.txt"
    txt_file.write_text("hello")
    with pytest.raises(ValueError, match="not a DOCX"):
        parse_docx(txt_file)


def test_parse_docx_valid(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("Hello World")
    doc.add_paragraph("This is a test document")
    docx_path = tmp_path / "test.docx"
    doc.save(str(docx_path))

    result = parse_docx(docx_path)
    assert "Hello World" in result
    assert "test document" in result


def test_parse_docx_bytes_valid():
    from docx import Document
    import io

    doc = Document()
    doc.add_paragraph("Test content from bytes")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    result = parse_docx_bytes(buf.getvalue())
    assert "Test content from bytes" in result


def test_parse_docx_empty(tmp_path):
    from docx import Document

    doc = Document()
    docx_path = tmp_path / "empty.docx"
    doc.save(str(docx_path))

    result = parse_docx(docx_path)
    assert result == ""
